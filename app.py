import io
import re
import csv
import zipfile
from datetime import datetime
from urllib.parse import urlparse
from zoneinfo import ZoneInfo

import streamlit as st
import requests
from bs4 import BeautifulSoup, Tag, NavigableString
from docx import Document
from docx.text.paragraph import Paragraph
from docx.oxml import OxmlElement

# -------------------------
# CONFIG / CONSTANTS
# -------------------------
ALWAYS_STRIP = {"script", "style", "noscript", "template"}
DEFAULT_EXCLUDE = [
    "header", "footer", "nav",
    ".cookie", ".newsletter",
    "[class*='breadcrumb']",
    "[class*='wishlist']",
    "[class*='simplesearch']",
    "[id*='gallery']",
    "[class*='usp']",
    "[class*='feefo']",
    "[class*='associated-blogs']",
    "[class*='popular']",
]
DATE_TZ = "Europe/London"
DATE_FMT = "%d/%m/%Y"  # UK format

# -------------------------
# UTILITIES
# -------------------------
def uk_today_str():
    return datetime.now(ZoneInfo(DATE_TZ)).strftime(DATE_FMT)

def clean_slug_to_name(slug: str) -> str:
    return slug.replace("-", " ").strip().title()

def fallback_page_name_from_url(url: str) -> str:
    path = urlparse(url).path.strip("/")
    parts = [p for p in path.split("/") if p]
    try:
        i = parts.index("destinations")
        if len(parts) > i + 2:
            return clean_slug_to_name(parts[i + 2])
    except ValueError:
        pass
    return clean_slug_to_name(parts[-1] if parts else (urlparse(url).hostname or "Page"))

def fetch_html(url: str) -> tuple[str, str]:
    resp = requests.get(url, timeout=30)
    resp.raise_for_status()
    return resp.url, resp.text

def normalise_keep_newlines(s: str) -> str:
    s = s.replace("\r\n", "\n").replace("\r", "\n").replace("\xa0", " ")
    s = re.sub(r"[ \t]+", " ", s)
    # keep explicit newlines; trim spaces around them
    s = re.sub(r"[ \t]*\n[ \t]*", "\n", s)
    return s

def annotate_anchor_text(a: Tag, annotate_links: bool) -> str:
    text = a.get_text(" ", strip=True)
    href = a.get("href", "")
    return f"{text} (→ {href})" if (annotate_links and href) else text

def extract_text_preserve_breaks(node: Tag | NavigableString, annotate_links: bool) -> str:
    """Extract visible text; convert <br> to '\n'; handle anchors as one unit."""
    if isinstance(node, NavigableString):
        return str(node)
    parts = []
    for child in node.children:
        if isinstance(child, NavigableString):
            parts.append(str(child))
        elif isinstance(child, Tag):
            if child.name == "br":
                parts.append("\n")
            elif child.name == "a":
                parts.append(annotate_anchor_text(child, annotate_links))
            else:
                parts.append(extract_text_preserve_breaks(child, annotate_links))
    return "".join(parts)

def extract_signposted_lines_from_body(body: Tag, annotate_links: bool) -> list[str]:
    """
    Emit ONLY:
      - <h1> … <h6> lines
      - <p> lines
    Lists flattened to <p>. Critically, <p> is split on <br> and blank lines preserved
    (blank <p> emitted as '<p>' with no text).
    """
    lines: list[str] = []

    def emit_lines(tag_name: str, text: str):
        text = normalise_keep_newlines(text)
        segments = text.split("\n")  # preserve blanks
        for seg in segments:
            seg_stripped = seg.strip()
            if seg_stripped:
                lines.append(f"<{tag_name}> {seg_stripped}")
            else:
                if tag_name == "p":
                    lines.append("<p>")  # explicit blank line
                # we typically don't emit blank headings

    def handle(tag: Tag):
        name = tag.name
        if name in ALWAYS_STRIP:
            return

        if name in {"h1","h2","h3","h4","h5","h6"}:
            txt = extract_text_preserve_breaks(tag, annotate_links)
            if txt.strip():
                emit_lines(name, txt)

        elif name == "p":
            txt = extract_text_preserve_breaks(tag, annotate_links)
            if txt.strip() or "\n" in txt:  # allow blank lines from <br><br>
                emit_lines("p", txt)

        elif name in {"ul", "ol"}:
            for li in tag.find_all("li", recursive=False):
                txt = extract_text_preserve_breaks(li, annotate_links)
                if txt.strip():
                    emit_lines("p", txt)
                # one nested level
                for sub in li.find_all(["ul", "ol"], recursive=False):
                    for sub_li in sub.find_all("li", recursive=False):
                        sub_txt = extract_text_preserve_breaks(sub_li, annotate_links)
                        if sub_txt.strip():
                            emit_lines("p", sub_txt)

        # ignore everything else

        # Recurse to capture nested text blocks
        for child in tag.children:
            if isinstance(child, Tag):
                handle(child)

    for child in body.children:
        if isinstance(child, Tag):
            handle(child)

    # Deduplicate trivial adjacent repeats
    deduped, prev = [], None
    for ln in lines:
        if ln != prev:
            deduped.append(ln)
        prev = ln
    return deduped

def first_h1_text(soup: BeautifulSoup) -> str | None:
    if not soup.body:
        return None
    h1 = soup.body.find("h1")
    if not h1:
        return None
    txt = extract_text_preserve_breaks(h1, annotate_links=False)
    txt = normalise_keep_newlines(txt)
    return txt.strip() or None

# -------------------------
# DOCX helpers
# -------------------------
def iter_paragraphs_and_tables(doc: Document):
    for p in doc.paragraphs:
        yield p
    for tbl in doc.tables:
        for row in tbl.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    yield p

def replace_placeholders_safe(doc: Document, mapping: dict[str, str]):
    """Replace placeholders safely: longer keys first to avoid nested/partial clashes."""
    keys = sorted(mapping.keys(), key=len, reverse=True)
    for p in iter_paragraphs_and_tables(doc):
        t = p.text or ""
        replaced = False
        for k in keys:
            v = mapping[k]
            if k in t:
                t = t.replace(k, v)
                replaced = True
        if replaced:
            # collapse runs (ok for placeholders)
            for r in list(p.runs):
                r.clear()
            p.clear()
            p.add_run(t)

def find_placeholder_paragraph(doc: Document, placeholder: str) -> Paragraph | None:
    for p in iter_paragraphs_and_tables(doc):
        if placeholder in (p.text or ""):
            return p
    return None

def insert_paragraph_after(paragraph: Paragraph, text: str = "") -> Paragraph:
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    new_para = Paragraph(new_p, paragraph._parent)
    if text:
        new_para.add_run(text)
    return new_para

def replace_placeholder_with_lines(doc: Document, placeholder: str, lines: list[str]):
    target = find_placeholder_paragraph(doc, placeholder)
    if target is None:
        raise ValueError(f"Placeholder '{placeholder}' not found in template.")
    if not lines:
        target.clear()
        return
    # first line replaces the placeholder, rest are new paragraphs after
    target.clear()
    target.add_run(lines[0])
    anchor = target
    for line in lines[1:]:
        anchor = insert_paragraph_after(anchor, line)

def build_docx(template_bytes: bytes, meta: dict, lines: list[str]) -> bytes:
    bio = io.BytesIO(template_bytes)
    doc = Document(bio)

    # Support both bracketed and bare DESCRIPTION tokens
    replace_placeholders_safe(doc, {
        "[PAGE]": meta["page"],
        "[DATE]": meta["date"],
        "[URL]": meta["url"],
        "[TITLE]": meta["title"],
        "[TITLE LENGTH]": str(meta["title_len"]),
        "[DESCRIPTION]": meta["description"],
        "DESCRIPTION": meta["description"],
        "[DESCRIPTION LENGTH]": str(meta["description_len"]),
    })

    replace_placeholder_with_lines(doc, "[PAGE BODY CONTENT]", lines)

    out = io.BytesIO()
    doc.save(out)
    out.seek(0)
    return out.read()

# -------------------------
# CORE PROCESS
# -------------------------
def process_url(
    url: str,
    exclude_selectors: list[str],
    annotate_links: bool = False,
):
    final_url, html = fetch_html(url)
    soup = BeautifulSoup(html, "lxml")

    # global strip
    for el in soup.find_all(list(ALWAYS_STRIP)):
        el.decompose()

    body = soup.body or soup

    # exclude universal blocks
    for sel in exclude_selectors:
        try:
            for el in body.select(sel):
                el.decompose()
        except Exception:
            pass

    # extract signposted lines
    lines = extract_signposted_lines_from_body(body, annotate_links=annotate_links)

    # meta
    head = soup.head or soup
    title = head.title.string.strip() if (head and head.title and head.title.string) else "N/A"
    meta_el = head.find("meta", attrs={"name": "description"}) if head else None
    description = meta_el.get("content").strip() if (meta_el and meta_el.get("content")) else "N/A"

    # page name: prefer H1
    page_name = first_h1_text(soup) or fallback_page_name_from_url(final_url)

    meta = {
        "page": page_name,
        "date": uk_today_str(),
        "url": final_url,
        "title": title,
        "title_len": len(title) if title != "N/A" else 0,
        "description": description,
        "description_len": len(description) if description != "N/A" else 0,
    }
    return meta, lines

# -------------------------
# STREAMLIT APP
# -------------------------
st.set_page_config(page_title="Explore Template Autofill", page_icon="364704cc-6899-4fc3-b37c-29dbfd0a4f3f.png", layout="wide")

st.markdown(
    """
    <style>
    @import url('https://fonts.googleapis.com/css2?family=Manrope:wght@400;600&display=swap');

    html, body, [class*="css"] {
        font-family: 'Manrope', sans-serif;
        background-color: #537DFC !important;
        color: #FFFFFF !important;
    }

    .stApp {
        background-color: #537DFC;
    }

    h1, h2, h3, h4, h5, h6 {
        color: #FFFFFF
    }

    .stTextInput > div > div > input,
    .stTextArea textarea,
    .stSelectbox > div > div,
    .stFileUploader > div,
    .stButton > button,
    .stDownloadButton > button,
    .stToggleSwitch,
    .stRadio > div > label,
    .stCheckbox > label {
        background-color: rgba(255, 255, 255, 0.1);
        color: #FFFFFF;
        border: 1px solid rgba(255, 255, 255, 0.3);
        border-radius: 6px;
    }

    .stButton > button:hover,
    .stDownloadButton > button:hover {
        background-color: rgba(255, 255, 255, 0.2);
    }

    .stTextInput > div > div > input::placeholder,
    .stTextArea textarea::placeholder {
        color: #DDDDDD;
    }

    .stSidebar, .st-cg, .st-c1, .css-1d391kg, .css-1kyxreq {
        background-color: rgba(255,255,255,0.1) !important;
        color: #FFFFFF !important;
        border-radius: 8px;
        padding: 10px;
    }

    .stTabs [role="tab"] {
        background-color: rgba(255,255,255,0.1);
        color: #FFFFFF;
        border: none;
        border-radius: 8px 8px 0 0;
        margin-right: 4px;
    }

    .stTabs [role="tab"][aria-selected="true"] {
        background-color: rgba(255,255,255,0.2);
        color: #FFFFFF;
        font-weight: bold;
    }

    .stDataFrame, .css-1cpxqw2 {
        background-color: rgba(255,255,255,0.05);
        border-radius: 6px;
    }

    .st-emotion-cache-1gv3huu eczjsme18 h2 h3 {
    color:#000000;
    }
    
    .css-1cpxqw2, .stDataFrame thead, .stDataFrame tbody {
        color: #FFFFFF;
    }

    .stCaption, .stExpanderHeader {
        color: #DDDDDD !important;
    }

    .stExpander {
        background-color: rgba(255, 255, 255, 0.05);
        border-radius: 6px;
    }

    .stDivider {
        border-color: rgba(255,255,255,0.3);
    }
    </style>
    """,
    unsafe_allow_html=True
)

st.title("Explore Template Autofill (Web)")

with st.sidebar:
    st.header("Template & Options")
    tpl_file = st.file_uploader("Upload Rec Template.docx", type=["docx"])
    st.caption("This should be your blank template with placeholders (e.g., [PAGE], [DATE], [PAGE BODY CONTENT], etc.).")

    st.divider()
    st.subheader("Exclude Selectors")
    exclude_txt = st.text_area(
        "Comma-separated CSS selectors to remove from <body>",
        value=", ".join(DEFAULT_EXCLUDE),
        height=120
    )
    exclude_selectors = [s.strip() for s in exclude_txt.split(",") if s.strip()]

    st.subheader("Link formatting")
    annotate_links = st.toggle("Append (→ URL) after anchor text", value=False)

    st.caption("Timezone fixed to Europe/London; dates in DD/MM/YYYY.")

tab1, tab2 = st.tabs(["Single URL", "Batch (CSV)"])

with tab1:
    st.subheader("Single page")
    url = st.text_input("URL", value="https://www.explore.co.uk/destinations/asia/cambodia")
    col_a, col_b = st.columns([1,1])
    with col_a:
        do_preview = st.button("Extract preview")
    with col_b:
        do_doc = st.button("Generate DOCX")

    if do_preview or do_doc:
        if not tpl_file and do_doc:
            st.error("Please upload your Rec Template.docx in the sidebar first.")
        else:
            try:
                meta, lines = process_url(url, exclude_selectors, annotate_links=annotate_links)
                st.success("Extracted successfully.")
                with st.expander("Meta (preview)", expanded=True):
                    st.write(meta)
                with st.expander("Signposted content (preview)", expanded=True):
                    st.text("\n".join(lines))

                if do_doc:
                    out_bytes = build_docx(tpl_file.read(), meta, lines)
                    st.download_button(
                        "Download DOCX",
                        data=out_bytes,
                        file_name=f"{meta['page']} - Content Recommendations.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    )
            except Exception as e:
                st.exception(e)

with tab2:
    st.subheader("Batch process CSV")
    st.caption("Upload a CSV with a header row; required column: url. Optional: out_name.")
    batch_file = st.file_uploader("CSV file", type=["csv"], key="csv")
    if st.button("Run batch"):
        if not tpl_file:
            st.error("Please upload your Rec Template.docx in the sidebar first.")
        elif not batch_file:
            st.error("Please upload a CSV.")
        else:
            tpl_bytes = tpl_file.read()
            rows = list(csv.DictReader(io.StringIO(batch_file.getvalue().decode("utf-8"))))
            if not rows:
                st.error("CSV appears empty.")
            elif "url" not in rows[0]:
                st.error("CSV must include a 'url' column.")
            else:
                memzip = io.BytesIO()
                zf = zipfile.ZipFile(memzip, "w", zipfile.ZIP_DEFLATED)
                results = []
                for i, row in enumerate(rows, 1):
                    u = row["url"].strip()
                    try:
                        meta, lines = process_url(u, exclude_selectors, annotate_links=annotate_links)
                        out_name = (row.get("out_name") or f"{meta['page']} - Content Recommendations").strip()
                        out_bytes = build_docx(tpl_bytes, meta, lines)
                        zf.writestr(f"{out_name}.docx", out_bytes)
                        results.append({"url": u, "status": "ok", "file": f"{out_name}.docx"})
                    except Exception as e:
                        results.append({"url": u, "status": f"error: {e}", "file": ""})
                zf.close()
                memzip.seek(0)
                st.success("Batch complete.")
                st.dataframe(results)
                st.download_button(
                    "Download ZIP",
                    data=memzip.read(),
                    file_name="content_recommendations.zip",
                    mime="application/zip",
                )
